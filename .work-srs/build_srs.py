from __future__ import annotations

import hashlib
import math
import re
import shutil
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "579126340-Tai-liệu-SRS-Mau-1.docx"
OUTPUT = ROOT / "SRS-He-thong-dat-ve-xe-khach-truc-tuyen.docx"
EXPECTED_REFERENCE_HASH = "548477a2408267f9d30323a42774e0e262e6bf2b4665ed892c043cffbdb523b0"
SKILL_SCRIPTS = Path(
    r"C:\Users\minht\.codex\plugins\cache\openai-primary-runtime\documents\26.805.11740\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = "2E74B5"
LIGHT_BLUE = "5B9BD5"
NAVY = "17365D"
PALE_BLUE = "DDEBF7"
PALEST_BLUE = "F3F7FB"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = "000000"
CONTENT_WIDTH_DXA = 9120
TABLE_INDENT_DXA = 120


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_rfonts(rpr, name: str) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_run_font(run, name="Times New Roman", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    set_rfonts(run._element.rPr, name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Times New Roman", size=11, color=BLACK, bold=None, italic=None):
    style.font.name = name
    style._element.get_or_add_rPr()
    set_rfonts(style._element.rPr, name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def ensure_style(doc, name: str, base="Normal"):
    if name in doc.styles:
        return doc.styles[name]
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base in doc.styles:
        style.base_style = doc.styles[base]
    return style


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        1: (15, BLUE, 18, 6),
        2: (13, LIGHT_BLUE, 14, 4),
        3: (11.5, NAVY, 10, 3),
        4: (11, BLACK, 8, 2),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style_name = f"Heading {level}"
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True
        style.paragraph_format.line_spacing = 1.05

    if "List Paragraph" in doc.styles:
        lp = doc.styles["List Paragraph"]
        set_style_font(lp, size=11)
        lp.paragraph_format.space_after = Pt(3)
        lp.paragraph_format.line_spacing = 1.1

    title_project = ensure_style(doc, "SRS Project Title")
    set_style_font(title_project, size=24, color=BLUE, bold=True)
    title_project.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_project.paragraph_format.space_after = Pt(14)
    title_project.paragraph_format.keep_together = True

    title_doc = ensure_style(doc, "SRS Document Title")
    set_style_font(title_doc, size=18, color=NAVY, bold=True)
    title_doc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_doc.paragraph_format.space_after = Pt(18)

    subtitle = ensure_style(doc, "SRS Subtitle")
    set_style_font(subtitle, size=11, color=GRAY, italic=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)

    fm = ensure_style(doc, "Front Matter Heading")
    set_style_font(fm, size=15, color=BLUE, bold=True)
    fm.paragraph_format.space_before = Pt(6)
    fm.paragraph_format.space_after = Pt(10)
    fm.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"] if "Caption" in doc.styles else ensure_style(doc, "Caption")
    set_style_font(caption, size=9, color=GRAY, italic=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.widow_control = True

    code = ensure_style(doc, "Code Block")
    set_style_font(code, name="Consolas", size=8.5, color="1F1F1F")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    note = ensure_style(doc, "SRS Note")
    set_style_font(note, size=10.5, color=NAVY)
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.right_indent = Inches(0.08)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.1

    small = ensure_style(doc, "SRS Small")
    set_style_font(small, size=9, color=GRAY)
    small.paragraph_format.space_after = Pt(3)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color="D9E2F3", size=4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{side}")
        edge = borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_field(paragraph, instruction: str, result: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run._r.append(begin)
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)
    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    if result:
        res_run = paragraph.add_run(result)
        set_run_font(res_run, size=10, color=GRAY)
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def create_abstract_numbering(doc: Document, kind: str, levels=4) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if kind == "bullet":
            lvl_text.set(qn("w:val"), ["•", "–", "▪", "◦"][level % 4])
        else:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(720 + level * 360))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 360))
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        if kind == "bullet":
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), "Arial")
            rfonts.set(qn("w:hAnsi"), "Arial")
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_num_instance(doc: Document, abstract_id: int, restart=True) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    if restart:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numpr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    nid = numpr.find(qn("w:numId"))
    if nid is None:
        nid = OxmlElement("w:numId")
        numpr.append(nid)
    nid.set(qn("w:val"), str(num_id))


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")


def clean_inline(text: str) -> str:
    return text.replace("<br>", " ").replace("<br/>", " ").strip()


def add_rich_text(paragraph, text: str, default_size=11, default_color=BLACK) -> None:
    text = clean_inline(text)
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, default_size - 1), color=NAVY)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEF3F8")
            run._r.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label = link.group(1) if link else token
            run = paragraph.add_run(label)
            set_run_font(run, size=default_size, color=BLUE)
            run.underline = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, color=default_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def shade_paragraph(paragraph, fill: str, left_border: str | None = None) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if left_border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "16")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        pbdr.append(left)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="SRS Note")
    shade_paragraph(p, PALE_BLUE, BLUE)
    add_rich_text(p, text, default_size=10.5, default_color=NAVY)


def add_code_block(doc: Document, lines: list[str], language: str) -> None:
    if language.lower() == "mermaid":
        add_note(doc, "Sơ đồ nguồn Mermaid đã được thay bằng hình trực quan đã kiểm tra ở mục sơ đồ liên quan của chương này.")
        return
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, LIGHT_GRAY, "B7B7B7")
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="1F1F1F")


def add_heading(doc: Document, text: str, level: int, heading_num_id: int, page_break=False):
    level = max(1, min(4, level))
    p = doc.add_paragraph(style=f"Heading {level}")
    add_rich_text(p, text, default_size={1: 15, 2: 13, 3: 11.5, 4: 11}[level], default_color={1: BLUE, 2: LIGHT_BLUE, 3: NAVY, 4: BLACK}[level])
    for run in p.runs:
        run.bold = True
    apply_num(p, heading_num_id, level - 1)
    set_outline_level(p, level)
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", text).strip()


def split_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    parts = re.split(r"(?<!\\)\|", line)
    return [p.strip().replace("\\|", "|") for p in parts]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", c.replace(" ", "")) for c in cells)


def col_weights(rows: list[list[str]]) -> list[float]:
    cols = len(rows[0])
    weights = []
    for idx in range(cols):
        texts = [re.sub(r"[*_`]", "", row[idx]) for row in rows if idx < len(row)]
        max_len = max((len(t) for t in texts), default=8)
        header = texts[0].lower() if texts else ""
        if header in {"id", "stt", "mức", "role", "bước"} or max_len <= 8:
            weight = 0.65
        elif max_len <= 18:
            weight = 1.0
        elif max_len <= 35:
            weight = 1.45
        else:
            weight = 2.1
        weights.append(weight)
    return weights


def add_seq_caption(paragraph, label: str, number: int, title: str) -> None:
    run = paragraph.add_run(f"{label} ")
    set_run_font(run, size=9, color=GRAY, bold=True, italic=True)
    add_field(paragraph, f"SEQ {label} \\* ARABIC", str(number))
    run = paragraph.add_run(f". {title}")
    set_run_font(run, size=9, color=GRAY, italic=True)


def add_markdown_table(doc: Document, rows: list[list[str]], title: str, table_index: int) -> None:
    if not rows or not rows[0]:
        return
    cols = len(rows[0])
    clean_rows = []
    for row in rows:
        row = row[:cols] + [""] * max(0, cols - len(row))
        clean_rows.append(row)

    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.keep_with_next = True
    add_seq_caption(cap, "Bảng", table_index, title)

    table = doc.add_table(rows=len(clean_rows), cols=cols)
    table.style = "Table Grid"
    weights = col_weights(clean_rows)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)

    font_size = 9.5 if cols <= 4 else 8.5
    for ri, row in enumerate(clean_rows):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_shading(cell, NAVY if ri == 0 else (PALEST_BLUE if ri % 2 == 0 else WHITE))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ci == 0 and len(text) < 18) or len(text) <= 3 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, text, default_size=font_size, default_color=WHITE if ri == 0 else BLACK)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
        keep_row_together(table.rows[ri])
    mark_repeat_header(table.rows[0])
    apply_table_geometry(table, widths, indent_dxa=TABLE_INDENT_DXA)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_list_paragraph(doc: Document, text: str, num_id: int, level: int) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_after = Pt(3)
    apply_num(p, num_id, min(level, 3))
    add_rich_text(p, text)


@dataclass
class ParseState:
    heading_num_id: int
    bullet_abstract_id: int
    number_abstract_id: int
    table_index: int = 0
    last_heading: str = "Nội dung đặc tả"


def add_markdown(doc: Document, path: Path, state: ParseState, heading_offset: int, skip_first_h1: bool) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    active_list_kind = None
    active_list_num = None
    first_h1_skipped = False

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(x.strip() for x in paragraph_buffer if x.strip())
            if text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.widow_control = True
                add_rich_text(p, text)
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(raw)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            active_list_kind = None
            active_list_num = None
            in_code = True
            code_lang = stripped[3:].strip()
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            active_list_kind = None
            active_list_num = None
            md_level = len(heading.group(1))
            text = strip_heading_number(heading.group(2))
            if md_level == 1 and skip_first_h1 and not first_h1_skipped:
                first_h1_skipped = True
            else:
                level = min(4, md_level + heading_offset)
                add_heading(doc, text, level, state.heading_num_id)
                state.last_heading = text
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1].strip()):
            flush_paragraph()
            active_list_kind = None
            active_list_num = None
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            state.table_index += 1
            add_markdown_table(doc, rows, state.last_heading, state.table_index)
            continue

        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.+)$", raw)
        if list_match:
            flush_paragraph()
            indent = max(0, len(list_match.group(1).replace("\t", "    ")) // 2)
            marker = list_match.group(2)
            kind = "number" if marker[0].isdigit() else "bullet"
            if kind != active_list_kind or active_list_num is None:
                abstract = state.number_abstract_id if kind == "number" else state.bullet_abstract_id
                active_list_num = new_num_instance(doc, abstract, restart=True)
                active_list_kind = kind
            add_list_paragraph(doc, list_match.group(3), active_list_num, indent)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            active_list_kind = None
            active_list_num = None
            add_note(doc, stripped.lstrip("> "))
            i += 1
            continue

        if not stripped or stripped == "---":
            flush_paragraph()
            active_list_kind = None
            active_list_num = None
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    if in_code and code_lines:
        add_code_block(doc, code_lines, code_lang)


def add_document_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=NAVY):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for ci, header in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ""
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_rich_text(p, header, default_size=9, default_color=WHITE)
        for run in p.runs:
            run.bold = True
    mark_repeat_header(table.rows[0])
    for ri, values in enumerate(rows, start=1):
        for ci, value in enumerate(values):
            cell = table.cell(ri, ci)
            cell.text = ""
            set_cell_shading(cell, PALEST_BLUE if ri % 2 == 0 else WHITE)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, value, default_size=9)
        keep_row_together(table.rows[ri])
    apply_table_geometry(table, widths, indent_dxa=TABLE_INDENT_DXA)
    return table


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PBL6 — TÀI LIỆU PHÂN TÍCH HỆ THỐNG")
    set_run_font(run, size=10, color=GRAY, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)

    doc.add_paragraph("HỆ THỐNG ĐẶT VÉ XE KHÁCH\nTRỰC TUYẾN", style="SRS Project Title")
    doc.add_paragraph("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)", style="SRS Document Title")
    doc.add_paragraph("Software Requirements Specification", style="SRS Subtitle")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("Từ tổng quan sản phẩm đến yêu cầu nghiệp vụ, kỹ thuật và nghiệm thu")
    set_run_font(run, size=11, color=NAVY, italic=True)

    meta = [
        ("Mã dự án", "PBL6"),
        ("Mã tài liệu", "PBL6-SRS-001"),
        ("Phiên bản", "1.0.0"),
        ("Trạng thái", "Bản đặc tả chi tiết"),
        ("Phân loại", "Nội bộ dự án"),
    ]
    table = doc.add_table(rows=len(meta) + 1, cols=2)
    table.style = "Table Grid"
    for ci, value in enumerate(("Thuộc tính", "Giá trị")):
        cell = table.cell(0, ci)
        cell.text = ""
        set_cell_borders(cell, color="D9E2F3", size=4)
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_rich_text(p, value, default_size=9, default_color=WHITE)
        for run in p.runs:
            run.bold = True
    mark_repeat_header(table.rows[0])
    for ri, (label, value) in enumerate(meta, start=1):
        left, right = table.rows[ri].cells
        for cell in (left, right):
            set_cell_borders(cell, color="D9E2F3", size=4)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, PALE_BLUE)
        set_cell_shading(right, WHITE)
        left.text = ""
        right.text = ""
        p1 = left.paragraphs[0]
        p2 = right.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        add_rich_text(p1, label, default_size=10, default_color=NAVY)
        for run in p1.runs:
            run.bold = True
        add_rich_text(p2, value, default_size=10)
    apply_table_geometry(table, [2400, 6720], indent_dxa=TABLE_INDENT_DXA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("Đà Nẵng, 20/08/2026")
    set_run_font(r, size=10, color=GRAY, italic=True)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_paragraph("PHÊ DUYỆT TÀI LIỆU", style="Front Matter Heading")
    add_document_table(
        doc,
        ["Vai trò", "Họ và tên / Chức danh", "Ký xác nhận"],
        [
            ["Người biên soạn", "Nhóm dự án PBL6", "…………………………………"],
            ["Người xem xét", "<Chưa xác định>", "…………………………………"],
            ["Người phê duyệt", "<Chưa xác định>", "…………………………………"],
        ],
        [2200, 4520, 2400],
    )
    doc.add_paragraph()
    doc.add_paragraph("LỊCH SỬ THAY ĐỔI", style="Front Matter Heading")
    add_document_table(
        doc,
        ["Ngày hiệu lực", "A/M/D", "Người cập nhật", "Mô tả thay đổi", "Phiên bản"],
        [["20/08/2026", "A", "Nhóm dự án PBL6", "Tạo mới tài liệu SRS đầy đủ từ bộ đặc tả mô-đun và thư viện sơ đồ", "1.0.0"]],
        [1300, 800, 1800, 4000, 1220],
    )
    p = doc.add_paragraph(style="SRS Small")
    p.add_run("Quy ước: A — Thêm; M — Sửa; D — Xóa.")
    doc.add_page_break()

    doc.add_paragraph("MỤC LỤC", style="Front Matter Heading")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Mục lục được cập nhật tự động khi mở tài liệu trong Microsoft Word.")
    doc.add_page_break()

    doc.add_paragraph("DANH MỤC HÌNH", style="Front Matter Heading")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Hình"', "Danh mục hình được cập nhật tự động khi mở tài liệu trong Microsoft Word.")
    doc.add_paragraph("DANH MỤC BẢNG", style="Front Matter Heading")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Bảng"', "Danh mục bảng được cập nhật tự động khi mở tài liệu trong Microsoft Word.")
    doc.add_page_break()


@dataclass(frozen=True)
class FigureSpec:
    path: str
    title: str
    description: str


FIGURE_GROUPS: dict[str, list[FigureSpec]] = {
    "overview": [
        FigureSpec("overview/system-context.png", "Bối cảnh hệ thống", "Thể hiện các nhóm người dùng, ba kênh client, ranh giới nền tảng và hai tích hợp bên ngoài chính."),
        FigureSpec("use-cases/use-cases-customer.png", "Use Case của Guest và Customer", "Tổng hợp mục tiêu tìm kiếm, giữ ghế, tạo booking, thanh toán, quản lý vé và đánh giá."),
        FigureSpec("use-cases/use-cases-operations.png", "Use Case vận hành nhà xe", "Mô tả phạm vi Operator Staff và Driver trong quản lý đội xe, chuyến, manifest và check-in."),
        FigureSpec("use-cases/use-cases-admin.png", "Use Case quản trị nền tảng", "Mô tả quyền quản lý tổ chức, tài khoản, nội dung, giao dịch và báo cáo toàn nền tảng."),
    ],
    "process": [
        FigureSpec("processes/activity-booking.png", "Activity Diagram quy trình đặt vé", "Diễn giải các bước và nhánh quyết định từ tìm chuyến đến phát hành vé hoặc phục hồi khi giao dịch thất bại."),
    ],
    "usecases": [
        FigureSpec("robustness/robustness-login.png", "Robustness Diagram đăng nhập", "Phân tách boundary, control và entity khi xác thực, kiểm tra tài khoản và phát hành phiên truy cập."),
        FigureSpec("robustness/robustness-booking.png", "Robustness Diagram đặt vé", "Làm rõ trách nhiệm giữa UI, API/control, SeatHold, Booking, Payment và Ticket."),
        FigureSpec("robustness/robustness-cancel-ticket.png", "Robustness Diagram hủy vé", "Thể hiện preview chính sách, xác nhận hủy, giải phóng ghế và yêu cầu hoàn tiền."),
        FigureSpec("robustness/robustness-create-trip.png", "Robustness Diagram tạo chuyến", "Mô tả kiểm tra tenant, cấu hình chuyến, xung đột xe/tài xế và phát hành Trip."),
        FigureSpec("robustness/robustness-check-in.png", "Robustness Diagram check-in", "Mô tả quét mã vé, xác minh quyền/chuyến và chuyển trạng thái Ticket an toàn."),
        FigureSpec("robustness/robustness-cancel-trip.png", "Robustness Diagram hủy chuyến", "Làm rõ tác động dây chuyền đến booking, ticket, refund, notification và reporting."),
        FigureSpec("sequences/sequence-seat-hold.png", "Sequence Diagram giữ ghế", "Chi tiết khóa, kiểm tra toàn bộ ghế, giao dịch all-or-nothing, TTL và idempotency."),
        FigureSpec("sequences/sequence-create-booking.png", "Sequence Diagram tạo booking", "Chi tiết xác minh hold, snapshot giá, passenger và tạo Booking chờ thanh toán."),
        FigureSpec("sequences/sequence-payment-provider.png", "Sequence Diagram tích hợp cổng thanh toán", "Thể hiện payment intent, chuyển hướng, callback/webhook và kiểm tra chữ ký."),
        FigureSpec("sequences/sequence-payment-confirm-booking.png", "Sequence Diagram xác nhận Booking sau thanh toán", "Mô tả phối hợp Payment–Booking để chốt ghế và phát hành Ticket."),
        FigureSpec("sequences/sequence-ticket-delivery.png", "Sequence Diagram phát hành và gửi vé", "Thể hiện tạo QR, yêu cầu thông báo và cơ chế retry không làm rollback booking đã commit."),
        FigureSpec("sequences/sequence-cancel-preview.png", "Sequence Diagram preview hủy", "Mô tả cách tính điều kiện, phí và số tiền hoàn trước khi Customer xác nhận."),
        FigureSpec("sequences/sequence-cancel-trip.png", "Sequence Diagram hủy chuyến", "Chi tiết phát sự kiện TripCancelled và xử lý các booking/ticket bị ảnh hưởng."),
        FigureSpec("sequences/sequence-refund-saga.png", "Sequence Diagram Refund Saga", "Mô tả xử lý hoàn tiền, retry, trạng thái lỗi và đối soát thủ công."),
        FigureSpec("sequences/sequence-publish-trip.png", "Sequence Diagram phát hành chuyến", "Thể hiện kiểm tra lịch, lưu Trip, outbox event và tạo inventory ghế ở Booking Service."),
    ],
    "states": [
        FigureSpec("states/state-trip-seat-hold.png", "State Diagram TripSeat và SeatHold", "Biểu diễn vòng đời AVAILABLE–HELD–BOOKED–DISABLED và quy tắc hết hạn hold."),
        FigureSpec("states/state-booking-payment.png", "State Diagram Booking và Payment", "Thể hiện sự phối hợp giữa trạng thái đơn đặt chỗ và kết quả thanh toán."),
        FigureSpec("states/state-ticket-refund.png", "State Diagram Ticket và Refund", "Làm rõ phát hành, check-in, sử dụng, hủy vé và các trạng thái hoàn tiền."),
        FigureSpec("states/state-trip.png", "State Diagram Trip", "Mô tả chuyển trạng thái chuyến từ nháp/phát hành đến hoàn tất hoặc hủy."),
    ],
    "data": [
        FigureSpec("domain-models/domain-identity.png", "Mô hình miền Identity", "Thể hiện User, role, membership, tenant và các đối tượng xác minh/xác thực."),
        FigureSpec("domain-models/domain-transport.png", "Mô hình miền Transport", "Thể hiện tổ chức, xe, sơ đồ ghế, tài xế, tuyến, điểm dừng và Trip."),
        FigureSpec("domain-models/domain-booking.png", "Mô hình miền Booking", "Thể hiện TripSnapshot, TripSeat, SeatHold, Booking, Passenger và Ticket."),
        FigureSpec("domain-models/domain-payment.png", "Mô hình miền Payment", "Thể hiện Payment, PaymentAttempt, WebhookReceipt, Refund và ReconciliationCase."),
        FigureSpec("data-models/erd-identity.png", "ERD Identity Service", "Cụ thể hóa khóa, quan hệ và phạm vi dữ liệu định danh/tenant."),
        FigureSpec("data-models/erd-transport.png", "ERD Transport Service", "Cụ thể hóa dữ liệu đội xe, tuyến, điểm dừng, sơ đồ ghế và chuyến."),
        FigureSpec("data-models/erd-booking.png", "ERD Booking Service", "Cụ thể hóa inventory ghế theo chuyến, hold, booking, passenger và ticket."),
        FigureSpec("data-models/erd-payment.png", "ERD Payment Service", "Cụ thể hóa giao dịch, lần thử, webhook, refund và đối soát."),
    ],
    "architecture": [
        FigureSpec("architecture/microservices-architecture.png", "Kiến trúc Microservices", "Thể hiện client, API Gateway, ranh giới service, database sở hữu riêng, Redis, broker và observability."),
        FigureSpec("events/event-payment.png", "Luồng sự kiện thanh toán", "Mô tả các event và consumer từ payment request đến xác nhận booking/ticket."),
        FigureSpec("events/event-cancellation.png", "Luồng sự kiện hủy và hoàn tiền", "Mô tả propagation của cancellation/refund tới booking, payment, notification và reporting."),
        FigureSpec("events/event-trip.png", "Luồng sự kiện chuyến", "Mô tả publish/update/cancel Trip và cách các service cập nhật snapshot/read model."),
        FigureSpec("deployment/deployment.png", "Mô hình triển khai mục tiêu", "Thể hiện routing, workload container, dữ liệu, message broker, storage và observability ở môi trường triển khai."),
        FigureSpec("deployment/deployment-local-demo.png", "Mô hình triển khai local/demo", "Mô tả baseline chạy toàn hệ thống bằng container cho phát triển, tích hợp và nghiệm thu."),
    ],
}


def add_figure(doc: Document, spec: FigureSpec, number: int) -> None:
    path = ROOT / ".qa" / "final-all" / spec.path
    if not path.exists():
        add_note(doc, f"Thiếu tài sản sơ đồ: {spec.path}")
        return
    intro = doc.add_paragraph()
    intro.paragraph_format.page_break_before = True
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("Mục đích sơ đồ. ")
    set_run_font(r, size=10, color=NAVY, bold=True)
    r = intro.add_run(spec.description)
    set_run_font(r, size=10, color=GRAY)

    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    max_w, max_h = 6.05, 6.85
    width = min(max_w, max_h * aspect)
    height = width / aspect
    if height > max_h:
        height = max_h
        width = height * aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    inline._inline.docPr.set("descr", spec.description)
    inline._inline.docPr.set("title", spec.title)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_seq_caption(cap, "Hình", number, spec.title)


def add_figure_group(doc: Document, group: str, state: ParseState, figure_index: int) -> int:
    specs = FIGURE_GROUPS[group]
    add_heading(doc, "Sơ đồ minh họa và đặc tả trực quan", 2, state.heading_num_id)
    add_note(doc, "Các sơ đồ dưới đây là bản diễn giải trực quan của yêu cầu có ID. Khi có mâu thuẫn, yêu cầu văn bản, Business Rule và Acceptance Criteria là nguồn quyết định.")
    for spec in specs:
        figure_index += 1
        add_figure(doc, spec, figure_index)
    return figure_index


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        section.different_first_page_header_footer = True


def configure_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("PBL6-SRS-001  |  HỆ THỐNG ĐẶT VÉ XE KHÁCH TRỰC TUYẾN")
        set_run_font(r, size=8, color=GRAY)
        ppr = hp._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "D9E2F3")
        pbdr.append(bottom)
        ppr.append(pbdr)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Phiên bản 1.0.0  •  Nội bộ dự án  |  Trang ")
        set_run_font(r, size=8, color=GRAY)
        add_field(fp, "PAGE", "1")
        r = fp.add_run(" / ")
        set_run_font(r, size=8, color=GRAY)
        add_field(fp, "NUMPAGES", "1")


def add_source_appendix(doc: Document, state: ParseState, heading_num_id: int, source_files: list[Path]) -> None:
    add_heading(doc, "PHỤ LỤC — NGUỒN TÀI LIỆU VÀ QUẢN LÝ THAY ĐỔI", 1, heading_num_id, page_break=True)
    add_heading(doc, "Nguồn nội dung", 2, heading_num_id)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_rich_text(p, "Bản SRS này hợp nhất nội dung mô-đun theo thứ tự từ tổng quan đến chi tiết. Đường dẫn dưới đây là nguồn duy trì; mọi thay đổi hành vi phải cập nhật yêu cầu, quy tắc, use case, contract và tiêu chí nghiệm thu liên quan.")
    state.table_index += 1
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.keep_with_next = True
    add_seq_caption(cap, "Bảng", state.table_index, "Danh mục nguồn nội dung SRS")
    rows = []
    for idx, path in enumerate(source_files, start=1):
        rows.append([str(idx), str(path.relative_to(ROOT)).replace("\\", "/"), "Nguồn đặc tả được hợp nhất nguyên nghĩa vào tài liệu này"])
    add_document_table(doc, ["STT", "Tệp nguồn", "Vai trò"], rows, [700, 4520, 3900])

    add_heading(doc, "Nguyên tắc quản lý thay đổi", 2, heading_num_id)
    principles = [
        "Thay đổi hành vi phải cập nhật Functional Requirement và Business Rule trước.",
        "Thay đổi luồng phải cập nhật Use Case, trạng thái liên quan và nhánh lỗi.",
        "Thay đổi contract phải cập nhật API/Event contract và quy tắc tương thích.",
        "Thay đổi chỉ hoàn tất khi Acceptance Criteria, test và ma trận truy vết đã được cập nhật.",
        "Diagram là diễn giải trực quan; ID yêu cầu và văn bản SRS là nguồn quyết định.",
    ]
    num = new_num_instance(doc, state.bullet_abstract_id)
    for item in principles:
        add_list_paragraph(doc, item, num, 0)

    add_heading(doc, "Thông tin cần phê duyệt trước baseline", 2, heading_num_id)
    add_note(doc, "Các trường người xem xét/người phê duyệt hiện để trống có chủ đích. Nhóm dự án cần điền tên, chức danh, ký xác nhận và chốt nhà cung cấp thanh toán/thông báo trước khi phát hành baseline chính thức.")


def build() -> None:
    if sha256(REFERENCE) != EXPECTED_REFERENCE_HASH:
        raise RuntimeError("Reference DOCX hash mismatch; fresh template distillation is required.")
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)
    configure_styles(doc)
    configure_page(doc)
    configure_header_footer(doc)
    set_update_fields(doc)

    props = doc.core_properties
    props.title = "SRS — Hệ thống đặt vé xe khách trực tuyến"
    props.subject = "Đặc tả yêu cầu phần mềm từ tổng quan đến chi tiết"
    props.author = "Nhóm dự án PBL6"
    props.keywords = "SRS, bus ticket, booking, payment, microservices, PBL6"
    props.comments = "Hợp nhất từ bộ tài liệu docs/srs và thư viện sơ đồ đã kiểm tra."
    props.version = "1.0.0"

    heading_abstract = create_abstract_numbering(doc, "number", levels=4)
    heading_num_id = new_num_instance(doc, heading_abstract, restart=True)
    bullet_abstract = create_abstract_numbering(doc, "bullet", levels=4)
    number_abstract = create_abstract_numbering(doc, "number", levels=4)
    state = ParseState(heading_num_id, bullet_abstract, number_abstract)

    add_cover(doc)
    add_front_matter(doc)

    chapters = [
        (
            "GIỚI THIỆU CHUNG",
            [(ROOT / "docs/srs/README.md", 0, True)],
            None,
        ),
        (
            "TỔNG QUAN SẢN PHẨM, PHẠM VI VÀ TÁC NHÂN",
            [
                (ROOT / "docs/srs/overview/product-overview.md", 1, False),
                (ROOT / "docs/srs/overview/system-scope-and-context.md", 1, False),
                (ROOT / "docs/srs/overview/actors-and-permissions.md", 1, False),
            ],
            "overview",
        ),
        (
            "QUY TRÌNH VÀ QUY TẮC NGHIỆP VỤ",
            [
                (ROOT / "docs/srs/business/business-processes.md", 1, False),
                (ROOT / "docs/srs/business/business-rules.md", 1, False),
            ],
            "process",
        ),
        (
            "ĐẶC TẢ USE CASE",
            [
                (ROOT / "docs/srs/business/use-cases/README.md", 1, False),
                (ROOT / "docs/srs/business/use-cases/identity-and-access.md", 1, False),
                (ROOT / "docs/srs/business/use-cases/journey-search-and-booking.md", 1, False),
                (ROOT / "docs/srs/business/use-cases/payment-cancellation-and-change.md", 1, False),
                (ROOT / "docs/srs/business/use-cases/trip-operations.md", 1, False),
                (ROOT / "docs/srs/business/use-cases/administration-and-reporting.md", 1, False),
            ],
            "usecases",
        ),
        (
            "ĐẶC TẢ YÊU CẦU CHỨC NĂNG",
            [(ROOT / "docs/srs/requirements/functional-requirements.md", 0, True)],
            None,
        ),
        (
            "YÊU CẦU VỀ TRẠNG THÁI NGHIỆP VỤ",
            [(ROOT / "docs/srs/requirements/state-requirements.md", 0, True)],
            "states",
        ),
        (
            "YÊU CẦU MIỀN NGHIỆP VỤ VÀ DỮ LIỆU",
            [(ROOT / "docs/srs/requirements/data-requirements.md", 0, True)],
            "data",
        ),
        (
            "YÊU CẦU GIAO DIỆN NGƯỜI DÙNG",
            [(ROOT / "docs/srs/requirements/user-interface-requirements.md", 0, True)],
            None,
        ),
        (
            "YÊU CẦU CHẤT LƯỢNG, BẢO MẬT VÀ VẬN HÀNH",
            [(ROOT / "docs/srs/requirements/quality-requirements.md", 0, True)],
            None,
        ),
        (
            "RÀNG BUỘC KIẾN TRÚC VÀ GIAO DIỆN DỊCH VỤ",
            [
                (ROOT / "docs/srs/architecture/service-architecture-constraints.md", 1, False),
                (ROOT / "docs/srs/architecture/service-interfaces.md", 1, False),
                (ROOT / "docs/srs/architecture/exceptions-and-recovery.md", 1, False),
            ],
            "architecture",
        ),
        (
            "NGHIỆM THU VÀ TRUY VẾT YÊU CẦU",
            [(ROOT / "docs/srs/verification/acceptance-and-traceability.md", 0, True)],
            None,
        ),
    ]

    source_files: list[Path] = []
    figure_index = 0
    for title, files, figure_group in chapters:
        add_heading(doc, title, 1, heading_num_id, page_break=True)
        for path, offset, skip_h1 in files:
            source_files.append(path)
            add_markdown(doc, path, state, offset, skip_h1)
        if figure_group:
            figure_index = add_figure_group(doc, figure_group, state, figure_index)

    add_source_appendix(doc, state, heading_num_id, source_files)

    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal":
            p.paragraph_format.widow_control = True
    doc.save(OUTPUT)

    if sha256(REFERENCE) != EXPECTED_REFERENCE_HASH:
        raise RuntimeError("Reference DOCX changed during authoring.")
    print(f"OUTPUT={OUTPUT}")
    print(f"REFERENCE_SHA256={sha256(REFERENCE)}")
    print(f"FIGURES={figure_index}")
    print(f"TABLES={state.table_index}")


if __name__ == "__main__":
    build()
