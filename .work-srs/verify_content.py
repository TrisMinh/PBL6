import re
from pathlib import Path

from docx import Document

root = Path(__file__).resolve().parents[1]
src_text = "\n".join(p.read_text(encoding="utf-8") for p in sorted((root / "docs/srs").rglob("*.md")))
doc = Document(root / "SRS-He-thong-dat-ve-xe-khach-truc-tuyen.docx")
parts = [p.text for p in doc.paragraphs]
for table in doc.tables:
    for row in table.rows:
        parts.extend(cell.text for cell in row.cells)
for section in doc.sections:
    parts.extend(p.text for p in section.header.paragraphs)
    parts.extend(p.text for p in section.footer.paragraphs)
doc_text = "\n".join(parts)
pattern = re.compile(r"\b(?:GOAL|BP|BR|UC|FR|NFR|AC|EVT|ERR)-[A-Z0-9][A-Z0-9.-]*\b")
source_ids = set(pattern.findall(src_text))
doc_ids = set(pattern.findall(doc_text))
missing = sorted(source_ids - doc_ids)
print("SOURCE_IDS", len(source_ids))
print("DOCUMENT_IDS", len(doc_ids))
print("MISSING", len(missing))
for item in missing:
    print(item)
raise SystemExit(1 if missing else 0)
