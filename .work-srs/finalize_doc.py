from docx import Document
from docx.oxml.ns import qn


PATH = "SRS-He-thong-dat-ve-xe-khach-truc-tuyen.docx"
doc = Document(PATH)

# The retained teaching template links Heading styles to a bullet list. Word
# normalizes direct numbering back to that style during field updates, which
# makes both headings and TOC entries display bullets. Remove the inherited
# list binding so hierarchy is carried by real Heading styles and TOC levels.
for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
    style = doc.styles[name]
    ppr = style._element.pPr
    if ppr is not None:
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            ppr.remove(num_pr)

for paragraph in doc.paragraphs:
    if paragraph.style and paragraph.style.name.startswith("Heading "):
        ppr = paragraph._p.pPr
        if ppr is not None:
            num_pr = ppr.find(qn("w:numPr"))
            if num_pr is not None:
                ppr.remove(num_pr)

# Strip only the bullet marker and its immediate spacing tab from already
# materialized TOC results. Keep hyperlinks and PAGEREF fields intact.
markers = {"•", "–", "▪", "◦"}
for paragraph in doc.paragraphs:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if not style_name.startswith("toc "):
        continue
    for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
        children = list(hyperlink)
        marker_index = None
        for idx, run in enumerate(children):
            texts = run.findall(qn("w:t"))
            if any((t.text or "") in markers for t in texts):
                marker_index = idx
                hyperlink.remove(run)
                break
        if marker_index is not None:
            remaining = list(hyperlink)
            if marker_index < len(remaining):
                candidate = remaining[marker_index]
                if candidate.find(qn("w:tab")) is not None:
                    rpr = candidate.find(qn("w:rPr"))
                    web_hidden = rpr.find(qn("w:webHidden")) if rpr is not None else None
                    if web_hidden is None:
                        hyperlink.remove(candidate)
        break

# Cached TOC/page values are now authoritative. Prevent an expensive automatic
# refresh on every open; users can still update fields manually when content changes.
settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is not None:
    update.set(qn("w:val"), "false")

doc.save(PATH)
print(PATH)
