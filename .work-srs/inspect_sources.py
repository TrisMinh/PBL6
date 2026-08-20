from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def val(obj):
    if obj is None:
        return None
    try:
        return float(obj.pt)
    except Exception:
        pass
    try:
        return float(obj.inches)
    except Exception:
        pass
    return str(obj)


def color(rgb):
    return str(rgb) if rgb is not None else None


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for idx, p in enumerate(doc.paragraphs):
        text = re.sub(r"\s+", " ", p.text).strip()
        if text:
            paragraphs.append(
                {
                    "index": idx,
                    "style": p.style.name if p.style else None,
                    "text": text,
                    "alignment": str(p.alignment),
                    "runs": [
                        {
                            "text": r.text,
                            "font": r.font.name,
                            "size_pt": val(r.font.size),
                            "bold": r.bold,
                            "italic": r.italic,
                            "color": color(r.font.color.rgb),
                        }
                        for r in p.runs
                        if r.text
                    ],
                }
            )

    styles = []
    for s in doc.styles:
        if s.type != 1:
            continue
        pf = s.paragraph_format
        styles.append(
            {
                "name": s.name,
                "base": s.base_style.name if s.base_style else None,
                "font": s.font.name,
                "size_pt": val(s.font.size),
                "bold": s.font.bold,
                "italic": s.font.italic,
                "color": color(s.font.color.rgb),
                "alignment": str(pf.alignment),
                "space_before_pt": val(pf.space_before),
                "space_after_pt": val(pf.space_after),
                "line_spacing": str(pf.line_spacing),
                "left_indent_in": val(pf.left_indent),
                "first_line_indent_in": val(pf.first_line_indent),
                "keep_with_next": pf.keep_with_next,
                "page_break_before": pf.page_break_before,
            }
        )

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([re.sub(r"\s+", " ", c.text).strip() for c in row.cells])
        tables.append({"index": ti, "rows": len(table.rows), "cols": len(table.columns), "text": rows})

    sections = []
    for s in doc.sections:
        sections.append(
            {
                "page_width_in": s.page_width.inches,
                "page_height_in": s.page_height.inches,
                "top_margin_in": s.top_margin.inches,
                "bottom_margin_in": s.bottom_margin.inches,
                "left_margin_in": s.left_margin.inches,
                "right_margin_in": s.right_margin.inches,
                "header_distance_in": s.header_distance.inches,
                "footer_distance_in": s.footer_distance.inches,
                "start_type": str(s.start_type),
                "header": [p.text for p in s.header.paragraphs],
                "footer": [p.text for p in s.footer.paragraphs],
            }
        )

    package = []
    with zipfile.ZipFile(path) as zf:
        for info in zf.infolist():
            package.append({"path": info.filename, "size": info.file_size, "crc": info.CRC})

    return {
        "path": str(path.resolve()),
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "paragraphs": paragraphs,
        "styles": styles,
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(doc.inline_shapes),
        "package": package,
    }


if __name__ == "__main__":
    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    out.write_text(json.dumps(inspect_docx(src), ensure_ascii=False, indent=2), encoding="utf-8")
    print(out)
